'''
Program:
	该程序可以协助识别网址中的不同部分，以便加快爬虫项目的开发。
History:
	2023/4/7	Shane	Second release
	finish function of createDoc
'''
from docx import Document
from docx.shared import Cm,Pt
from docx.document import Document as Doc
import os

def openDoc(docName):
    '''打开word文档'''
    path = os.getcwd()
    docPath = path + f'\\{docName}'
    print (docPath)
    #空循环，检测到有word文档为止
    while not os.path.exists (docPath):
        pass
    try:
        os.system(f'{docPath}')
    except:
        print('打开word文档失败')

def createDoc(message1, message2, docName):
    '''创建一个word文档，并将信息导入'''
    document = Document()       #创建文件对象
    document.add_paragraph(message1, style='List Bullet')   #添加无序列表
    document.add_paragraph (message2 , style = 'List Bullet')
    document.save (f'{docName}.docx')     #保存为

def tools_checkDif(url1, url2):
    '''对比两个网址，并输出其网址差异的地方'''
    firstNum, endNum = 0, 0
    saveNum = 2
    modUrl1 = list (url1)
    modUrl2 = list (url2)

    while modUrl1[firstNum] == modUrl2[firstNum]:
        firstNum += 1
    while modUrl1 [-endNum-1] == modUrl2 [-endNum-1]:
        endNum += 1
    #删除首端、末端相同值，并保留saveNum个数
    del modUrl1[0:firstNum - saveNum]
    del modUrl2 [0:firstNum - saveNum]
    if endNum != 0 and endNum>=saveNum:
        del modUrl1 [-endNum+saveNum : ]    #     del modUrl2 [-endNum+saveNum : ]
        del modUrl2 [-endNum+saveNum : ]    #     del modUrl2 [-endNum+saveNum : ]

    #打印修改后的列表
    modUrl1 = "".join (modUrl1)
    modUrl2 = "".join (modUrl2)
    print (modUrl1)
    print (modUrl2)
    return modUrl1, modUrl2

if __name__ == '__main__':
    # url1 = 'happy2happy'
    # url2 = 'happ23happy'
    url1 = input("请输入网址1：")
    url2 = input("请输入网址2：")
    url1, url2 = tools_checkDif(url1, url2)
    docName = 'demo.docx'
    createDoc(url1, url2, docName)
    openDoc (docName)
